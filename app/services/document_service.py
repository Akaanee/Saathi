import os
import io
import logging
from typing import Dict, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

logger = logging.getLogger(__name__)

class DocumentService:
    def __init__(self):
        self.output_dir = os.getenv("OUTPUT_DIR", "outputs")
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_legal_notice(
        self,
        complaint_data: Dict[str, Any],
        output_format: str = "docx"
    ) -> bytes:
        if output_format.lower() == "docx":
            return self._generate_docx_notice(complaint_data)
        elif output_format.lower() == "pdf":
            return self._generate_pdf_notice(complaint_data)
        else:
            raise ValueError(f"Unsupported format: {output_format}")

    def _generate_docx_notice(self, complaint_data: Dict[str, Any]) -> bytes:
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        title = doc.add_heading('LEGAL NOTICE', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")
        doc.add_paragraph()
        
        header = doc.add_heading('To,', level=1)
        
        respondent_info = complaint_data.get('respondent', {})
        doc.add_paragraph(f"Name: {respondent_info.get('name', '[Respondent Name]')}")
        doc.add_paragraph(f"Address: {respondent_info.get('address', '[Address]')}")
        doc.add_paragraph()
        
        subject = doc.add_heading('Subject: Legal Notice Under [Applicable Law]', level=1)
        subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        doc.add_paragraph("Dear Sir/Madam,")
        doc.add_paragraph()
        
        doc.add_heading('1. FACTS OF THE MATTER', level=2)
        factual_matrix = complaint_data.get('factual_matrix', complaint_data.get('incident_description', ''))
        doc.add_paragraph(f"""
{complaint_data.get('complainant_name', '[Complainant Name]')} 
(hereinafter referred to as the "Complainant") 
states as follows:
        """.strip())
        
        doc.add_paragraph(factual_matrix)
        doc.add_paragraph()
        
        doc.add_heading('2. GROUNDS OF COMPLAINT', level=2)
        legal_grounds = complaint_data.get('legal_grounds', '')
        if isinstance(legal_grounds, list):
            legal_grounds = '\n'.join([f"• {ground}" for ground in legal_grounds])
        doc.add_paragraph(legal_grounds)
        doc.add_paragraph()
        
        doc.add_heading('3. RELIEF SOUGHT', level=2)
        relief_sought = complaint_data.get('relief_sought', [])
        if isinstance(relief_sought, list):
            for i, relief in enumerate(relief_sought, 1):
                doc.add_paragraph(f"{i}. {relief}", style='List Number')
        else:
            doc.add_paragraph(relief_sought)
        doc.add_paragraph()
        
        doc.add_heading('4. NOTICE', level=2)
        doc.add_paragraph("""
You are hereby called upon to remedy the above-mentioned grievances within 15 days 
of receipt of this notice, failing which the Complainant shall be constrained to initiate 
appropriate legal proceedings before the competent court/forum without further reference 
to you, and in such event, you shall be solely liable for the costs and consequences thereof.
        """.strip())
        doc.add_paragraph()
        
        doc.add_paragraph("Yours faithfully,")
        doc.add_paragraph()
        doc.add_paragraph(f"Date: {datetime.now().strftime('%d-%m-%Y')}")
        doc.add_paragraph(f"Place: {complaint_data.get('incident_location', '[Location]')}")
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph(f"{complaint_data.get('complainant_name', '[Complainant Name]')}")
        doc.add_paragraph("[Signature]")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()

    def _generate_pdf_notice(self, complaint_data: Dict[str, Any]) -> bytes:
        pdf = FPDF()
        pdf.add_page()
        
        pdf.set_font('Helvetica', 'B', 16)
        pdf.cell(0, 10, 'LEGAL NOTICE', ln=True, align='C')
        pdf.ln(10)
        
        pdf.set_font('Helvetica', '', 12)
        pdf.cell(0, 8, f"Date: {datetime.now().strftime('%d-%m-%Y')}", ln=True)
        pdf.ln(5)
        
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, 'To,', ln=True)
        pdf.set_font('Helvetica', '', 12)
        
        respondent_info = complaint_data.get('respondent', {})
        pdf.cell(0, 8, f"Name: {respondent_info.get('name', '[Respondent Name]')}", ln=True)
        pdf.cell(0, 8, f"Address: {respondent_info.get('address', '[Address]')}", ln=True)
        pdf.ln(5)
        
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, 'Subject: Legal Notice', ln=True, align='C')
        pdf.ln(10)
        
        pdf.set_font('Helvetica', '', 12)
        pdf.multi_cell(0, 6, "Dear Sir/Madam,")
        pdf.ln(5)
        
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, '1. FACTS OF THE MATTER', ln=True)
        pdf.set_font('Helvetica', '', 12)
        
        factual_matrix = complaint_data.get('factual_matrix', complaint_data.get('incident_description', ''))
        pdf.multi_cell(0, 6, factual_matrix)
        pdf.ln(5)
        
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, '2. RELIEF SOUGHT', ln=True)
        pdf.set_font('Helvetica', '', 12)
        
        relief_sought = complaint_data.get('relief_sought', [])
        if isinstance(relief_sought, list):
            for relief in relief_sought:
                pdf.cell(0, 6, f"• {relief}", ln=True)
        else:
            pdf.multi_cell(0, 6, str(relief_sought))
        pdf.ln(5)
        
        pdf.set_font('Helvetica', 'B', 12)
        pdf.cell(0, 8, '3. NOTICE', ln=True)
        pdf.set_font('Helvetica', '', 12)
        pdf.multi_cell(0, 6, 
            "You are hereby called upon to remedy the above-mentioned grievances within "
            "15 days of receipt of this notice, failing which appropriate legal proceedings "
            "will be initiated."
        )
        pdf.ln(10)
        
        pdf.cell(0, 8, 'Yours faithfully,', ln=True)
        pdf.ln(10)
        pdf.cell(0, 8, f"Date: {datetime.now().strftime('%d-%m-%Y')}", ln=True)
        pdf.cell(0, 8, f"Place: {complaint_data.get('incident_location', '[Location]')}", ln=True)
        pdf.ln(15)
        pdf.cell(0, 8, f"{complaint_data.get('complainant_name', '[Complainant Name]')}", ln=True)
        
        return bytes(pdf.output())

    def generate_case_summary(
        self,
        complaint_data: Dict[str, Any],
        format: str = "txt"
    ) -> bytes:
        if format == "txt":
            return self._generate_txt_summary(complaint_data)
        elif format == "docx":
            return self._generate_docx_summary(complaint_data)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_txt_summary(self, complaint_data: Dict[str, Any]) -> bytes:
        summary_lines = []
        
        summary_lines.append("=" * 80)
        summary_lines.append("CASE SUMMARY")
        summary_lines.append("=" * 80)
        summary_lines.append("")
        
        summary_lines.append(f"Case ID: {complaint_data.get('case_id', 'N/A')}")
        summary_lines.append(f"Generated: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}")
        summary_lines.append("")
        
        summary_lines.append("-" * 80)
        summary_lines.append("PARTIES")
        summary_lines.append("-" * 80)
        summary_lines.append(f"Complainant: {complaint_data.get('complainant_name', 'N/A')}")
        summary_lines.append(f"Address: {complaint_data.get('complainant_address', 'N/A')}")
        summary_lines.append(f"Occupation: {complaint_data.get('complainant_occupation', 'N/A')}")
        summary_lines.append("")
        summary_lines.append(f"Respondent: {complaint_data.get('respondent', {}).get('name', 'N/A')}")
        summary_lines.append(f"Address: {complaint_data.get('respondent', {}).get('address', 'N/A')}")
        summary_lines.append("")
        
        summary_lines.append("-" * 80)
        summary_lines.append("KEY FACTS")
        summary_lines.append("-" * 80)
        summary_lines.append(f"Incident Date: {complaint_data.get('incident_date', 'N/A')}")
        summary_lines.append(f"Location: {complaint_data.get('incident_location', 'N/A')}")
        summary_lines.append("")
        summary_lines.append(complaint_data.get('incident_description', 'N/A'))
        summary_lines.append("")
        
        summary_lines.append("-" * 80)
        summary_lines.append("APPLICABLE LAWS")
        summary_lines.append("-" * 80)
        applicable_laws = complaint_data.get('applicable_laws', [])
        if isinstance(applicable_laws, list):
            for law in applicable_laws:
                summary_lines.append(f"• {law}")
        else:
            summary_lines.append(str(applicable_laws))
        summary_lines.append("")
        
        summary_lines.append("-" * 80)
        summary_lines.append("RELIEF SOUGHT")
        summary_lines.append("-" * 80)
        relief_sought = complaint_data.get('relief_sought', [])
        if isinstance(relief_sought, list):
            for relief in relief_sought:
                summary_lines.append(f"• {relief}")
        else:
            summary_lines.append(str(relief_sought))
        summary_lines.append("")
        
        summary_lines.append("=" * 80)
        summary_lines.append("END OF SUMMARY")
        summary_lines.append("=" * 80)
        
        content = "\n".join(summary_lines)
        return content.encode('utf-8')

    def _generate_docx_summary(self, complaint_data: Dict[str, Any]) -> bytes:
        doc = Document()
        
        title = doc.add_heading('Case Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Case ID: {complaint_data.get('case_id', 'N/A')}")
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%d-%m-%Y %H:%M:%S')}")
        doc.add_paragraph()
        
        doc.add_heading('Parties', level=1)
        doc.add_paragraph(f"Complainant: {complaint_data.get('complainant_name', 'N/A')}")
        doc.add_paragraph(f"Address: {complaint_data.get('complainant_address', 'N/A')}")
        doc.add_paragraph(f"Respondent: {complaint_data.get('respondent', {}).get('name', 'N/A')}")
        doc.add_paragraph()
        
        doc.add_heading('Key Facts', level=1)
        doc.add_paragraph(complaint_data.get('incident_description', 'N/A'))
        doc.add_paragraph()
        
        doc.add_heading('Applicable Laws', level=1)
        applicable_laws = complaint_data.get('applicable_laws', [])
        if isinstance(applicable_laws, list):
            for law in applicable_laws:
                doc.add_paragraph(f"• {law}")
        doc.add_paragraph()
        
        doc.add_heading('Relief Sought', level=1)
        relief_sought = complaint_data.get('relief_sought', [])
        if isinstance(relief_sought, list):
            for relief in relief_sought:
                doc.add_paragraph(f"• {relief}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()

    def save_document(
        self,
        content: bytes,
        filename: str,
        subdirectory: Optional[str] = None
    ) -> str:
        if subdirectory:
            save_dir = os.path.join(self.output_dir, subdirectory)
        else:
            save_dir = self.output_dir
        
        os.makedirs(save_dir, exist_ok=True)
        
        filepath = os.path.join(save_dir, filename)
        
        with open(filepath, 'wb') as f:
            f.write(content)
        
        logger.info(f"Document saved: {filepath}")
        return filepath

    def get_output_path(self, filename: str) -> str:
        return os.path.join(self.output_dir, filename)


document_service = DocumentService()
